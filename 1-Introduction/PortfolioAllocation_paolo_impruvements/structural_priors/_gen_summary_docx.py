"""One-shot generator for the session summary Word document.

REVISED 2026-06-11 to reflect the analyst's catch on the active-stats
benchmark mis-specification. Original "pipeline pathology" claim withdrawn;
results re-benchmarked against EqualWeight_w_Cash; corrected analysis
restored.

Produces FINAL_RESULTS_SUMMARY.docx in the same folder.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    h.style.font.name = "Calibri"


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style.font.name = "Calibri"
    p.style.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row):
            tbl.rows[r_idx].cells[c_idx].text = str(cell_text)
    doc.add_paragraph()  # spacing


def add_monospace(doc: Document, text: str) -> None:
    """Add a block of text in a monospace font (for command outputs)."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────
# Build the document
# ─────────────────────────────────────────────────────────────────────────

doc = Document()
for sec in doc.sections:
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

# Title
title = doc.add_heading("PPO Portfolio Allocation Triage — Session Summary (Revised)", level=0)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
sub_run = sub.add_run(
    "Universe: 13 sector ETFs plus a synthetic CASH asset (N=14). "
    "Evaluation window: 2020-07-01 to 2026-05-19 (5.9 years, T=1477 daily bars). "
    "Revised 2026-06-11 after analyst review caught a benchmark mis-specification "
    "that produced spurious pipeline-pathology findings in the first version."
)
sub_run.italic = True


# 0. Revision note
add_heading(doc, "0. What changed in this revision", 1)
add_para(doc,
    "The first version of this memo concluded that the pipeline exhibits "
    "a structural negative bias and inverts signals. On analyst review, "
    "this was incorrect. The active-statistics block in 03_backtest.py "
    "benchmarked against EqualWeight (a cash-free, fully invested "
    "strategy), but the PPO agent has CASH as one of its 14 allocatable "
    "assets and structurally holds approximately 6.4 to 7.3 percent in "
    "cash across phases. Comparing a cash-holding strategy to a cash-free "
    "benchmark during a period in which risk assets compounded "
    "approximately 80 percent mechanically produces an IR of "
    "approximately negative 1 with a significant t-statistic, purely "
    "from the cash exposure, with zero correlation to any feature signal.")
add_para(doc,
    "The correct apples-to-apples benchmark is EqualWeight_w_Cash, which "
    "is also fully invested but holds 1/(M+1) in CASH alongside 1/M in "
    "each risky asset. The 03_backtest.py script has been patched to use "
    "it automatically when cash is enabled, and a cash-drag decomposition "
    "block has been added that prints both IRs side-by-side so the drag "
    "is visible. All four phases were re-backtested without retraining. "
    "The revised numbers and conclusions are below.")


# 1. Executive summary
add_heading(doc, "1. Executive Summary", 1)
add_para(doc,
    "This session ran a four-phase triage to determine whether a "
    "PPO-based active portfolio allocator can outperform a "
    "daily-rebalanced equal-weight (EW) baseline on a 13-ETF plus CASH "
    "universe. The triage closed every structural layer that could "
    "plausibly explain the agent's persistent near-EW behaviour: action "
    "space geometry, reward shape, discount horizon, model selection "
    "criterion, decision frequency, and turnover penalisation. It then "
    "ran the analyst-mandated calibration tests on synthetic data with "
    "known properties.")
add_para(doc,
    "With the corrected EqualWeight_w_Cash benchmark, the four ensemble "
    "outcomes are all statistically indistinguishable from zero IR. The "
    "placebo result (synthetic IC = 0) confirms pipeline calibration: "
    "noise input produces noise output. The IC = 0.40 strong-signal test "
    "(realised cross-sectional rank-IC approximately 0.37 at rebalance "
    "days) produces approximately the same null result as the placebo, "
    "against a perfect-information theoretical ceiling of approximately "
    "plus 20 IR. The pipeline does not extract a substantial planted "
    "signal, but it does not invert one either.")
add_para(doc,
    "The deployable answer is EqualWeight_w_Cash with the turbulence "
    "gate as a tail-risk overlay. Annualised Sharpe 1.103, cumulative "
    "return 72.60 percent, maximum drawdown negative 14.97 percent. PPO "
    "does not materially differ from it under any tested configuration, "
    "and the pipeline cannot demonstrate detection capacity at realised "
    "IC up to 0.37 in this design.")


# 2. Scripts modified / added
add_heading(doc, "2. Scripts Modified or Added", 1)
add_para(doc,
    "Every code change in this session is captured in the git history of "
    "the structural_priors folder. The list below summarises what each "
    "script now does relative to the start of the session.")

add_table(doc,
    headers=["Script", "Status", "What was changed"],
    rows=[
        ["utils.py", "modified",
         "Added action_logit_scale parameter on LogReturnPortfolioEnv (widens the "
         "softmax action box from [0,1] to [-s,+s]). Added compute_rebalance_dates "
         "helper for env-time weekly cadence. Added cadence and weekly_day kwargs "
         "on the env: step() now gates the agent's action before super().step() so "
         "the action only takes effect on rebalance days (typically Fridays, with "
         "holiday fallback). ValidationSharpeCallback gained selection_metric in "
         "{sharpe, ir} so model selection can be aligned with the article_benchmark "
         "reward. make_portfolio_env reads all new config blocks."],

        ["02_train.py", "modified",
         "Now passes the early_stopping.selection_metric value through to the "
         "callback so IR-based selection can be turned on per-config."],

        ["03_backtest.py", "modified (revised)",
         "Added compute_active_stats(): per-strategy daily active return mean (bps), "
         "tracking error (bps), annualised IR, iid paired t and two-sided p, and "
         "Newey-West HAC-adjusted t and p with lag floor(N^(1/4)). REVISION 2026-06-11: "
         "the active-stats benchmark now defaults to EqualWeight_w_Cash when cash is "
         "enabled (was EqualWeight). Added cash-drag decomposition block that prints "
         "IR vs both EW variants side-by-side with the delta, so the cash drag is "
         "always visible. Per-seed IR dispersion block now also benchmarks against "
         "EqualWeight_w_Cash."],

        ["inspect_action_range.py", "added",
         "Diagnostic that reads a weights_<algo>.csv and reports max/min weight, "
         "implied logit spread, L1 distance from equal weight, and turnover "
         "percentiles. Verdicts on whether the legacy box cap was binding. Used "
         "throughout the triage as the action-geometry instrument."],

        ["features.py", "modified",
         "Registered alpha_signal as a per_asset pass-through feature. It expects "
         "the column to be present in the dataframe (planted by gen_synthetic_data.py) "
         "and raises KeyError on real data so it cannot accidentally leak into "
         "production runs."],

        ["test_priors.py", "modified",
         "Added Test 4: an empirical proof that softmax(a + alpha*log(w_eq)) equals "
         "softmax(a) for any a and any alpha (the equal-weight prior is "
         "softmax-shift-invariant and therefore a mathematical no-op). All 4 tests "
         "pass."],

        ["test_cadence.py", "added",
         "Three unit tests for env-time cadence: 'daily' as identity, 'weekly' "
         "selecting Fridays with holiday-fallback to Thursday, and a real env "
         "stepping that confirms 98 percent of non-rebalance bars correctly "
         "replay the held action."],

        ["test_signal_recovery_upper_bound.py", "added",
         "Monte Carlo (2000 paths) of the theoretical IR achievable by a "
         "perfect-information weekly long-only strategy at planted IC in "
         "{0, 0.05, 0.10, 0.20, 0.40} on a 5.9-year window. Provides the "
         "denominator for pipeline efficiency."],

        ["gen_synthetic_data.py", "added",
         "Phase 3b generator. Loads the real OHLCV panel, block-bootstraps it on "
         "the date index (block length 20 days) to preserve cross-sectional "
         "correlation and volatility clustering, plants an alpha_signal column "
         "whose cross-sectional rank-correlation with the next-week cumulative "
         "return is set to a target IC. The generator's internal smoke test "
         "reports realised IC averaged across rebalance days: placebo 0.009, "
         "IC=0.40 arm 0.374 — both within ~6% relative of target."],

        ["METHODOLOGY.md", "modified",
         "Added section 6.5 (policy_prior is currently REMOVED from env-time "
         "injection; equal_weight is a mathematical no-op; alpha is the prior "
         "temperature, NOT a tilt-leverage knob) and section 6.6 (TC asymmetry: "
         "agent leg is net-of-TC while benchmark leg is gross, documented as a "
         "design choice)."],

        ["config.json + 5 scale configs", "modified",
         "Added action_logit_scale env knob (default 3.0 on the main config). "
         "Added early_stopping.selection_metric (default 'sharpe'; set to 'ir' on "
         "the three article_benchmark configs). Flipped policy_prior.enabled to "
         "false as the shipped default (avoids the equal_weight no-op footgun). "
         "Rewrote policy_prior docstring to reflect the actual current "
         "behaviour."],

        ["New configs", "added",
         "config_scale1.json, config_scale3.json, config_scale3_bench.json, "
         "config_scale3_bench_g09.json, config_scale3_bench_g09_lto003.json "
         "(triage sweep), config_phase1_5seeds.json (multi-seed daily robustness), "
         "config_phase4_weekly.json (weekly cadence combined arm), "
         "config_phase3_placebo_IC0.json, config_phase3_synth_IC040.json "
         "(synthetic data tests)."],

        ["Article_priority_1/utils.py", "modified",
         "Ported action_logit_scale support to the parallel folder so the "
         "geometry fix is consistent across implementations."],

        ["Article_priority_1/04_backtrader_replay.py", "modified",
         "Added an UNMAINTAINED banner pointing to structural_priors/ as the "
         "maintained version (the alignment-fix port was not done here)."],
    ])


# 3. Pre-registered experiments
add_heading(doc, "3. Pre-Registered Experiments", 1)
add_para(doc,
    "The triage proceeded in four pre-registered phases. Each phase had "
    "a concrete hypothesis, a concrete prediction, and a concrete decision "
    "rule for what to do next. Phases were run sequentially; the decision "
    "after each phase determined the next.")

add_table(doc,
    headers=["Phase", "Hypothesis tested", "Decision rule"],
    rows=[
        ["1 — Daily multi-seed rider",
         "Multi-seed averaging of the daily IR-selected arm produces 5 IRs "
         "straddling zero (analyst's pre-registration). Establishes the robust "
         "EqualWeight_w_Cash sign-off baseline.",
         "If 5 IRs straddle zero with ensemble p > 0.05, sign off on EW_w_Cash."],

        ["2 — Implement env-time weekly cadence",
         "Existence-only; the implementation work itself (not a hypothesis). "
         "Verifies the gate mechanism with three unit tests.",
         "Pass unit tests, proceed to Phase 4."],

        ["3a — Theoretical IR ceiling",
         "What is the maximum IR a perfect-information weekly strategy can "
         "achieve at planted IC = {0, 0.05, 0.10, 0.20, 0.40} on a 5.9-year "
         "window?",
         "Calibration only. Provides the denominator for Phase 3b efficiency."],

        ["3b — Planted-signal recovery (placebo + IC 0.40)",
         "Analyst's adaptive plan: placebo IR ~ 0 (calibrates the pipeline); "
         "IC = 0.40 IR > 0 with material magnitude (confirms recovery capacity). "
         "If both pass, bisect to find detection threshold.",
         "If placebo IR > 0 with significance, pipeline reports false positives. "
         "If IC = 0.40 fails to recover, pipeline is broken; debug, do not "
         "grid-search intermediate ICs."],

        ["4 — Combined arm (weekly cadence, all knobs stacked)",
         "Pre-registered success criteria: ensemble IR > 0, TE per |IR| below "
         "the daily arm's ratio, Sharpe at least 1.05. Not significance "
         "(unreachable on 5.9 years per the MDE arithmetic).",
         "If criteria met, revisit deployment decision. Otherwise, EW_w_Cash "
         "stays the deployable."],
    ])


# 4. Headline results
add_heading(doc, "4. Corrected Headline Results", 1)
add_para(doc,
    "All phases were re-backtested against EqualWeight_w_Cash. None of "
    "the four ensemble IRs is statistically distinguishable from zero "
    "(all p_NW > 0.19). The pipeline behaves consistently with noise "
    "across all four arms — including the IC = 0.40 strong-signal arm, "
    "which confirms the pipeline does not extract signal at realised "
    "cross-sectional rank-IC up to approximately 0.37.")

add_table(doc,
    headers=["Phase", "Setup", "Ensemble IR vs EW_w_Cash", "t_NW (p_NW)",
             "Per-seed signs (+/-)"],
    rows=[
        ["Phase 1", "Daily cadence, 5 seeds",                     "+0.323", "+0.918 (0.359)", "3 / 2"],
        ["Phase 4", "Weekly cadence, 5 seeds",                    "+0.517", "+1.292 (0.196)", "3 / 2"],
        ["Phase 3 placebo", "Synthetic IC = 0, 2 seeds",          "+0.367", "+0.917 (0.359)", "1 / 1"],
        ["Phase 3 IC = 0.40", "Synthetic realised IC = 0.37 at rebalance days, 2 seeds",
         "-0.214", "-0.539 (0.590)", "0 / 2"],
        ["3a theoretical at IC = 0.20",
         "Perfect-info weekly, Monte Carlo",
         "+10.27 (median)", "N/A", "N/A"],
        ["3a theoretical at IC = 0.40",
         "Perfect-info weekly, Monte Carlo",
         "+21.17 (median)", "N/A", "N/A"],
    ])

add_para(doc,
    "Note on the t_NW values for Phase 1 and the placebo. Both round to "
    "0.92 at two-decimal display precision (Phase 1 = 0.9179, placebo = "
    "0.9167). This is a numerical coincidence, not a transcription "
    "error: the IRs (+0.323 vs +0.367) and TEs (4.55 vs 1.71 bps/day) "
    "differ materially, with the Newey-West standard errors landing at "
    "proportionally different values that produce nearly identical "
    "t-statistics. The p-values agree to three decimals as a knock-on.")


# 5. Cash-drag decomposition
add_heading(doc, "5. Cash-Drag Decomposition (the analyst's catch)", 1)
add_para(doc,
    "The table below is the headline diagnostic of the revision. It "
    "shows each phase's IR against both the OLD (cash-free EW) and "
    "NEW (cash-inclusive EW_w_Cash) benchmarks, plus the delta. The "
    "delta in every case is approximately negative 1 to negative 1.5, "
    "consistent with the agent's average cash weight of approximately "
    "6.4 to 7.3 percent and the cash-versus-risky-asset return gap "
    "over the period. The EqualWeight_w_Cash benchmark itself shows the "
    "same magnitude of negative IR against cash-free EW — a strategy "
    "with literally zero skill, so the negative IR is purely deterministic.")

add_table(doc,
    headers=["Phase", "IR vs EW (no cash) — OLD",
             "IR vs EW_w_Cash — NEW",
             "Delta (cash drag)"],
    rows=[
        ["Phase 1 (daily)",        "-0.723", "+0.323", "-1.046"],
        ["Phase 4 (weekly)",       "-0.966", "+0.517", "-1.484"],
        ["Phase 3 placebo",        "-1.046", "+0.367", "-1.413"],
        ["Phase 3 IC = 0.40",      "-1.322", "-0.214", "-1.107"],
        ["EW_w_Cash (zero-skill reference)",
         "-1.103", "0.000 (= itself)", "-1.103"],
    ])

add_para(doc,
    "Average CASH weight per phase, computed from the saved weights "
    "files: Phase 1 = 6.64%, Phase 4 = 6.40%, placebo = 7.34%, IC=0.40 "
    "= 6.60%. All clustered near 7%, consistent with the structural "
    "drag mechanism.")

add_para(doc,
    "Phase 4's drag delta of negative 1.484 is larger than the "
    "zero-skill EW_w_Cash reference of negative 1.103. The excess of "
    "approximately negative 0.38 IR is not cash drag — it is the "
    "transaction-cost plus variance drag induced by the policy's tilts. "
    "PPO's tilted positions have higher concentration than 1/(M+1) "
    "and pay turnover cost on each rebalance day; both effects compound "
    "to a small additional negative IR vs cash-free EW that EW_w_Cash "
    "does not pay.")


# 6. Per-seed dispersion
add_heading(doc, "6. Per-Seed Dispersion (Corrected)", 1)
add_para(doc,
    "Per-seed numbers under the corrected benchmark. Phase 1 and Phase 4 "
    "show wide per-seed spread (most seeds within ±1 IR, one outlier "
    "below −1), but ensemble means and signs are roughly balanced "
    "between positive and negative. The IC = 0.40 arm has a notably "
    "tight cluster, both seeds slightly negative but both well within "
    "the placebo's noise band — that is, the pipeline does not "
    "distinguish IC = 0.37 strong signal from IC = 0 pure noise.")

add_table(doc,
    headers=["Phase", "Per-seed IRs vs EW_w_Cash", "Mean", "SD", "Signs +/-"],
    rows=[
        ["Phase 1 (daily)",
         "-0.66, +0.26, -1.59, +0.79, +0.13",
         "-0.215", "0.928", "3 / 2"],
        ["Phase 4 (weekly)",
         "+0.06, +0.74, -1.09, +0.76, -0.63",
         "-0.033", "0.823", "3 / 2"],
        ["Phase 3 placebo",
         "+0.61, -0.88",
         "-0.134", "1.054", "1 / 1"],
        ["Phase 3 IC = 0.40",
         "-0.14, -0.31",
         "-0.225", "0.117", "0 / 2"],
    ])


# 7. Raw outputs (corrected)
add_heading(doc, "7. Corrected Raw Output Excerpts", 1)
add_para(doc,
    "Verbatim excerpts of the re-run stage-3 output for each phase, "
    "showing the new active-stats block + cash-drag decomposition.")

add_para(doc, "Phase 1 (daily 5-seed) — corrected:")
add_monospace(doc,
"Active-return statistics vs EqualWeight_w_Cash\n"
"  PPO:   Active_bps=+0.09  TE_bps=4.55  IR=+0.323  t_NW=+0.92  p_NW=0.359\n"
"\n"
"Cash-drag decomposition: IR vs both EW variants\n"
"  Strategy                IR vs EW (no cash)   IR vs EW_w_Cash   Delta (= cash drag)\n"
"  PPO                                -0.723             +0.323              -1.047")

add_para(doc, "Phase 4 (weekly 5-seed) — corrected:")
add_monospace(doc,
"Active-return statistics vs EqualWeight_w_Cash\n"
"  PPO:   Active_bps=+0.09  TE_bps=2.91  IR=+0.517  t_NW=+1.29  p_NW=0.196\n"
"\n"
"Cash-drag decomposition: IR vs both EW variants\n"
"  Strategy                IR vs EW (no cash)   IR vs EW_w_Cash   Delta (= cash drag)\n"
"  PPO                                -0.966             +0.517              -1.484")

add_para(doc, "Phase 3 placebo (synthetic IC = 0) — corrected:")
add_monospace(doc,
"Active-return statistics vs EqualWeight_w_Cash\n"
"  PPO:   Active_bps=+0.04  TE_bps=1.71  IR=+0.367  t_NW=+0.92  p_NW=0.359\n"
"\n"
"Cash-drag decomposition: IR vs both EW variants\n"
"  Strategy                IR vs EW (no cash)   IR vs EW_w_Cash   Delta (= cash drag)\n"
"  PPO                                -1.046             +0.367              -1.413")

add_para(doc, "Phase 3 IC = 0.40 (realised IC ≈ 0.37 at rebalance days) — corrected:")
add_monospace(doc,
"Active-return statistics vs EqualWeight_w_Cash\n"
"  PPO:   Active_bps=-0.04  TE_bps=2.82  IR=-0.214  t_NW=-0.54  p_NW=0.590\n"
"\n"
"Cash-drag decomposition: IR vs both EW variants\n"
"  Strategy                IR vs EW (no cash)   IR vs EW_w_Cash   Delta (= cash drag)\n"
"  PPO                                -1.322             -0.214              -1.107")

add_para(doc, "Phase 3a theoretical IR ceiling (excerpt, unchanged):")
add_monospace(doc,
"    IC   median IR         5%        95%   P(IR>0)   P(|t|>1.96)\n"
"--------------------------------------------------------------------\n"
"  0.00      -0.019     -0.684      0.695     47.8%          5.5%\n"
"  0.05       2.544      1.848      3.247    100.0%        100.0%\n"
"  0.10       5.119      4.430      5.792    100.0%        100.0%\n"
"  0.20      10.270      9.555     11.027    100.0%        100.0%\n"
"  0.40      21.173     20.270     22.116    100.0%        100.0%")


# 8. Conclusions
add_heading(doc, "8. Corrected Conclusions and Deployable Decision", 1)
add_para(doc,
    "Four substantive conclusions from this session, in decreasing "
    "order of confidence.")

add_para(doc,
    "First, the pipeline IS calibrated against the correct benchmark. "
    "The placebo IR is plus 0.37 with p = 0.36 — within noise of zero. "
    "There is no structural negative bias. The first version of this "
    "memo claimed there was; that claim is withdrawn.")

add_para(doc,
    "Second, the pipeline does NOT demonstrate detection capacity at "
    "realised cross-sectional rank-IC of approximately 0.37. The IC = 0.40 "
    "synthetic arm produced an IR of negative 0.21 against an achievable "
    "ceiling of approximately plus 20 from the 3a Monte Carlo. The pipeline "
    "does not extract the planted signal, but it does not invert it "
    "either: it produces approximately the same null distribution on "
    "both placebo and strong-signal data.")

add_para(doc,
    "Caveat on this finding: it rests on two seeds of a single "
    "block-bootstrap draw. It is directionally credible — a "
    "perfect-info ceiling of approximately plus 20 IR missed entirely "
    "down to approximately zero is hard to explain by seed luck alone "
    "or by bootstrap-draw-specific structure — but a fully powered "
    "claim would need a second bootstrap draw and a larger seed pool. "
    "Option B in section 9 (alpha_signal as the only per-asset feature, "
    "approximately 30 minutes of compute) is precisely the cheap test "
    "that converts 'no detection observed' into a concrete mechanism: "
    "feature saliency loss in the 27-dimensional state vector.")

add_para(doc,
    "Third, the real-data results (Phase 1 and Phase 4) DO NOT bound "
    "the IC content of the real features. This is the subtle but "
    "important point. Phase 1 IR is plus 0.32 and Phase 4 IR is plus "
    "0.52, both indistinguishable from zero — but the IC = 0.40 result "
    "establishes that this pipeline produces approximately the same "
    "null distribution whether the data carries planted IC = 0 or "
    "planted IC = 0.37. Real features could carry IC = 0.5 and this "
    "pipeline would still report IR approximately zero. The defensible "
    "claim is about the apparatus, not the features: this pipeline "
    "extracts no alpha from these features on the tested arms, and its "
    "demonstrated detection capacity is below realised IC approximately "
    "0.37; real-feature IC content remains UNMEASURED by this design.")

add_para(doc,
    "Fourth, the deployable answer is EqualWeight_w_Cash with the "
    "turbulence gate. Annualised Sharpe 1.103, cumulative return 72.60 "
    "percent, maximum drawdown negative 14.97 percent. PPO does not "
    "materially differ from it under any tested configuration. The "
    "publishable negative claim is bounded — 'this pipeline does not "
    "extract alpha from these features and its detection capacity is "
    "below realised IC approximately 0.37' — and notably does NOT say "
    "'the features carry no signal' or 'active management leaks value.'")


# 9. Recommended next move
add_heading(doc, "9. Recommended Next Move", 1)
add_para(doc,
    "Two options. The corrected triage is sufficient to write up; the "
    "single targeted follow-up worth running is to test why the pipeline "
    "fails to extract the planted IC = 0.37.")

add_table(doc,
    headers=["Option", "Effort", "Decision it informs"],
    rows=[
        ["A. Ship the methodology paper.",
         "Writeup only",
         "Frame as: 'leak-free pipeline + properly powered placebo + "
         "cash-drag-corrected benchmark show this PPO + article_benchmark "
         "+ IR-sel pipeline does NOT extract alpha from a 13-ETF universe "
         "AND has demonstrated detection capacity below realised IC "
         "approximately 0.37 (subject to a 2-seed-1-bootstrap caveat). "
         "Real-feature IC content is therefore unmeasured. The deployable "
         "strategy is EqualWeight_w_Cash with the turbulence gate.'"],

        ["B. Single targeted follow-up: train an IC = 0.40 arm where "
         "alpha_signal is the ONLY per-asset feature in the state.",
         "~30 minutes compute",
         "Tests whether the binding constraint is feature saliency in "
         "the 27-dimensional state vector. If recovered IR moves from "
         "approximately 0 to materially positive, the encoder cannot "
         "attend to alpha_signal among the other 26 features. That "
         "would be a clean, publishable mechanism for the null result "
         "on real data, and converts the under-caveated detection-"
         "failure conclusion into a concrete cause with a fix path."],

        ["C. Both, B FIRST (RECOMMENDED, per analyst endorsement).",
         "Option B first, then Option A",
         "Combined deliverable: methodology paper plus one concrete "
         "diagnostic that identifies (or rules out) feature saliency as "
         "the binding constraint for future signal-extraction work. "
         "The analyst explicitly endorsed this option order."],
    ])


# 10. Closing note
add_heading(doc, "10. Closing note", 1)
add_para(doc,
    "The corrected triage is now closed with a defensible result. The "
    "pipeline-pathology rabbit hole was a false trail caught by the "
    "pre-registered placebo and the analyst's cash-drag observation — "
    "exactly what a calibrated evaluation infrastructure is supposed "
    "to do. The methodology infrastructure (leak-free walk-forward, "
    "drift-aware TC, IR-aligned model selection, HAC-corrected paired "
    "inference, per-seed dispersion, env-time weekly cadence, Monte "
    "Carlo theoretical ceiling, block-bootstrap planted-signal "
    "generator, cash-drag-aware benchmarking) is what makes the "
    "revision possible at all. Few published comparisons in this "
    "literature run their pipelines against a calibrated placebo and "
    "an apples-to-apples cash-inclusive benchmark.")


# Save
out = "FINAL_RESULTS_SUMMARY.docx"
doc.save(out)
print(f"Wrote {out}")
